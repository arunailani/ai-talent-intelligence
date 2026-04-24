import os
import fitz  # pymupdf for PDF
import docx  # python-docx for Word


def extract_text(file_path: str) -> str:
    """
    Universal document extractor.
    Detects file type and routes to correct parser.

    Supported formats:
    - PDF (.pdf)  → pymupdf
    - Word (.docx) → python-docx
    - Text (.txt)  → plain read

    Returns clean extracted text as a string.
    """
    extension = os.path.splitext(file_path)[1].lower()

    if extension == ".pdf":
        return _extract_pdf(file_path)
    elif extension == ".docx":
        return _extract_docx(file_path)
    elif extension == ".txt":
        return _extract_txt(file_path)
    else:
        raise ValueError(
            f"Unsupported file format: {extension}. "
            f"Supported: .pdf, .docx, .txt"
        )


def _extract_pdf(file_path: str) -> str:
    """
    Extracts text from PDF using pymupdf.
    Handles Canva, Word-exported, and standard PDFs.
    """
    doc = fitz.open(file_path)
    full_text = ""
    for page in doc:
        full_text += page.get_text()
    doc.close()

    return _clean_text(full_text)


def _extract_docx(file_path: str) -> str:
    """
    Extracts text from Word .docx files.
    Reads all paragraphs and table cells.
    """
    doc = docx.Document(file_path)
    full_text = ""

    # Extract paragraphs
    for para in doc.paragraphs:
        if para.text.strip():
            full_text += para.text + "\n"

    # Extract text from tables too
    # Many resumes use tables for layout
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    full_text += cell.text + "\n"

    return _clean_text(full_text)


def _extract_txt(file_path: str) -> str:
    """
    Extracts text from plain .txt files.
    """
    with open(file_path, "r", encoding="utf-8") as f:
        full_text = f.read()

    return _clean_text(full_text)


def _clean_text(text: str) -> str:
    """
    Removes excessive blank lines and strips whitespace.
    Applied to all extracted text regardless of source format.
    """
    lines = [line.strip() for line in text.splitlines()]
    cleaned = "\n".join(line for line in lines if line)
    return cleaned